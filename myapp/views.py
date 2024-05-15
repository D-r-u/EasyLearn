from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse
from django.shortcuts import render
from yake import yake

from .models import *

# Create your views here.
def login(request):
    return render(request,'login_index.html')
def login_post(request):
    username=request.POST['textfield']
    password=request.POST['textfield2']
    var=Login.objects.filter(username=username,password=password)
    if var.exists():
        var = Login.objects.get(username=username, password=password)
        request.session["lid"]=var.id
        if var.type == "admin":
            return HttpResponse('''<script>alert("Logged in Successfully!");window.location="/myapp/adm_home/"</script>''')
        elif var.type == "staff":
            return HttpResponse('''<script>alert("Logged in Successfully!");window.location="/myapp/stf_home/"</script>''')
        elif var.type == "user":
            return HttpResponse('''<script>alert("Logged in Successfully!");window.location="/myapp/usr_home/"</script>''')
        else :
            return HttpResponse('''<script>alert("Log in Unsuccessful!");window.location="/myapp/login/"</script>''')
    else :
        return HttpResponse('''<script>alert("USER Does not exist");window.location="/myapp/login/"</script>''')

def adm_logout(request):
    request.session["lid"]=""
    return HttpResponse('''<script>alert("You have Logged out");window.location="/myapp/login/"</script>''')

def adm_home(request):
    return render(request, 'admin/adm_home_index.html')

def adm_sendreply(request,id):
    return render(request,'Admin/send_reply.html',{'id':id})
def adm_sendreply_post(request):
    reply=request.POST['reply']
    cid=request.POST['cid']
    obj=Complaints.objects.get(id=cid)
    obj.reply=reply
    obj.status='replied'
    obj.save()
    return HttpResponse('''<script>alert("Reply Sent");window.location="/myapp/view_complaints/"</script>''')

def adm_app_rej_staff(request):
    res = Staff.objects.filter(status='Pending')
    return render(request, 'Admin/apprve_rjct_staff.html', {'data':res})
def adm_app_rej_staff_post(request):
    search = request.POST['srch']
    res=Staff.objects.filter(name__icontains=search,status="Pending")
    return render(request, 'Admin/apprve_rjct_staff.html', {'data':res})

def approve_post(request,id):
    res=Staff.objects.filter(id=id).update(status="approved")
    return HttpResponse('''<script>alert("Approved!");window.location="/myapp/apprvd_nd_rjctd_staff/"</script>''')

def reject_post(request,id):
    res=Staff.objects.filter(id=id).update(status="rejected")
    return HttpResponse('''<script>alert("Rejected!");window.location="/myapp/apprvd_nd_rjctd_staff/"</script>''')

def adm_chngpasswd(request):
    return render(request,'Admin/Chng_psswrd.html')
def adm_chngpasswd_post(request):
    old_password=request.POST['textfield']
    new_password=request.POST['textfield2']
    confirm_password = request.POST['textfield3']
    obj=Login.objects.filter(id=request.session['lid'],password=old_password)
    if obj.exists():
        obj = Login.objects.get(id=request.session['lid'], password=old_password)

        obj.password=confirm_password
        obj.save()
        return HttpResponse(
            '''<script>alert("Password Changed Successfully.");window.location="/myapp/login/"</script>''')
    else :
        return HttpResponse('''<script>alert("Incorrect Password");window.location="/myapp/Change_password/"</script>''')
    old_password = request.POST['textfield']
    new_password = request.POST['textfield2']
    confirm_password = request.POST['textfield3']
    obj = Login.objects.filter(id=request.session['lid'], password=old_password)
    if obj.exists():
        obj = Login.objects.get(id=request.session['lid'], password=old_password)

        if (new_password == confirm_password):
            obj.password = confirm_password
            obj.save()
            return HttpResponse(
                '''<script>alert("Password Changed Successfully.");window.location="/myapp/login/"</script>''')
        else:
            return HttpResponse(
                '''<script>alert("New and confirm password must be same");window.location="/myapp/Change_password/"</script>''')
    else:
        return HttpResponse(
            '''<script>alert("Incorrect Password");window.location="/myapp/Change_password/"</script>''')

def adm_vw_rejected_staff(request):
    res = Staff.objects.filter(status='rejected')
    return render(request, 'Admin/rejected_staff.html', {'data':res})
def adm_vw_rejected_staff_post(request):
    search=request.POST['srch']
    res = Staff.objects.filter(name__icontains=search, status="rejected")
    return render(request, 'Admin/rejected_staff.html', {'data':res})

def adm_vw_apprej_staff(request):
    res = Staff.objects.filter(status='approved')
    return render(request, 'Admin/adm_vw_approved.html', {'data':res})
def adm_vw_apprej_staff_post(request):
    search=request.POST['srch']
    res=Staff.objects.filter(name__icontains=search,status="approved")
    return render(request, 'Admin/adm_vw_approved.html', {'data':res})

def adm_vw_complaints(request):
    res=Complaints.objects.all()
    return render(request,'Admin/view_complaints.html',{'data':res})
def adm_vw_complaints_post(request):
    from_date=request.POST['from_date']
    to_date=request.POST['to_date']
    res=Complaints.objects.filter(date__range=[from_date,to_date])
    return render(request,'Admin/view_complaints.html',{'data':res})

def adm_vw_staffcomplaints(request,id):
    res=Staff_Complaints.objects.filter(STAFF_id=id)
    return render(request,'Admin/view_staff_complaints.html',{'data':res,'id':id})
def adm_vw_staffcomplaints_post(request):
    sid=request.POST['sid']
    print(sid)
    from_date=request.POST['from_date']
    to_date=request.POST['to_date']
    res=Staff_Complaints.objects.filter(date__range=[from_date,to_date],STAFF_id=sid)
    return render(request,'Admin/view_staff_complaints.html',{'data':res,id:id})

def adm_send_staffcomp_reply(request,id):
    return render(request,'Admin/view_staff_complaints.html',{'id':id})
def adm_send_staffcomp_reply_post(request):
    reply=request.POST['reply']
    cid=request.POST['cid']
    obj=Complaints.objects.get(id=cid)
    obj.reply=reply
    obj.status='replied'
    obj.save()
    return HttpResponse('''<script>alert("Reply Sent");window.location="/myapp/view_complaints_staff/"</script>''')

def adm_vw_review(request):
    res=Reviews.objects.all()
    return render(request,'Admin/view_reviews.html',{'data':res})
def adm_vw_review_post(request):
    from_date=request.POST['from']
    to_date=request.POST['to']
    res=Reviews.objects.filter(date__range=[from_date,to_date])
    return render(request,'Admin/view_reviews.html',{'data':res})

def adm_vw_user(request):
    res=User.objects.all()
    return render(request,'Admin/view_user.html',{'data':res})
def adm_vw_user_post(request):
    srch=request.POST['textfield']
    res=User.objects.filter(name__icontains=srch)
    return render(request,'Admin/view_user.html',{'data':res})




#=========Staff=============================
def stf_home(request):
    return render(request,'staff/stf_home.html')

def stf_adquest(request):
    res = Test.objects.filter(STAFF__LOGIN_id=request.session['lid'])
    q = Question.objects.filter(LOGIN_id=request.session['lid'])
    return render(request,'staff/add_questions.html',{'data':res,'q':q})
def stf_adquest_post(request):
    qut = request.POST['quest']
    qut2 = request.POST.get('questions','')
    tq = TestQuestions()

    if 'quest' in request.POST:
        tq.question = qut
        tq.type = 'typed'
    else:
        tq.question = Question.objects.filter(id=qut2).question
        tq.type = 'list'

    tq.TEST_id = request.POST['tname']
    tq.option1 = request.POST['opt1']
    tq.option2 = request.POST['opt2']
    tq.option3 = request.POST['opt3']
    tq.option4 = request.POST['opt4']
    tq.correct_answer = request.POST['crctans']
    tq.save()
    return HttpResponse('''<script>alert("Question added successfully.");window.location="/myapp/stf_add_questions/"</script>''')


def stf_addtest(request):
    return render(request,'staff/add_test.html')
def stf_addtest_post(request):
    t_name=request.POST['tname']
    date=request.POST['date']
    time=request.POST['time']
    obj=Test()
    obj.testname=t_name
    obj.date=date
    obj.time=time
    obj.STAFF= Staff.objects.get(LOGIN_id=request.session['lid'])
    obj.save()
    return HttpResponse('''<script>alert("Test added successfully.");window.location="/myapp/stf_add_test/"</script>''')


def stf_appreview(request):
    return render(request,'staff/appreview.html')
def stf_appreview_post(request):
    review=request.POST['review']
    rating=request.POST['rating']
    res=Reviews()
    res.review=review
    res.rating=rating
    from datetime import datetime
    res.date=datetime.now().date()
    res.type='staff'
    res.LOGIN_id=request.session['lid']
    res.save()
    return HttpResponse('''<script>alert("Review has been sent.");window.location="/myapp/stf_home/"</script>''')



def stf_changepswd(request):
    return render(request,'staff/change_passwd.html')
def stf_changepswd_post(request):
    old_password=request.POST['textfield']
    new_password=request.POST['textfield2']
    confirm_password = request.POST['textfield3']
    obj = Login.objects.filter(id=request.session['lid'], password=old_password)
    if obj.exists():
        obj = Login.objects.get(id=request.session['lid'], password=old_password)

        if (new_password == confirm_password):
            obj.password = confirm_password
            obj.save()
            return HttpResponse(
                '''<script>alert("Password Changed Successfully.");window.location="/myapp/login/"</script>''')
        else:
            return HttpResponse(
                '''<script>alert("New and confirm password must be same");window.location="/myapp/stf_change_passwd/"</script>''')
    else:
        return HttpResponse(
            '''<script>alert("Incorrect Password");window.location="/myapp/stf_change_passwd/"</script>''')


def stf_delete_questions(request,id):
    TestQuestions.objects.filter(id=id).delete()
    return HttpResponse('''<script>alert("Question Deleted.");window.location="/myapp/stf_manage_test_questions/"</script>''')

def stf_delete_test(request,id):
    Test.objects.filter(id=id).delete()
    return HttpResponse('''<script>alert("Test Deleted.");window.location="/myapp/stf_manage_test/"</script>''')


def stf_edit_profile(request):
    res = Staff.objects.get(LOGIN=request.session['lid'])
    return render(request,'staff/edit_profile.html',{'staff':res})
def stf_edit_profile_post(request):
    name=request.POST['textfield']
    gender=request.POST['gender']
    dob=request.POST['textfield1']
    username=request.POST['textfield2']
    phone=request.POST['textfield3']
    mailid=request.POST['textfield4']
    state=request.POST['textfield5']
    house=request.POST['textfield6']
    place=request.POST['textfield7']
    district=request.POST['textfield8']
    pincode=request.POST['textfield9']
    qualification=request.POST['textfield10']
    current_inst=request.POST['textfield11']
    if 'imageField' in request.FILES:
        imageField = request.FILES['imageField']
        if imageField!="":
            from datetime import datetime
            date1 = "stf_photos/" + datetime.now().strftime("%Y%m%d-%H%M%S") + "-1.jpg"
            fs1 = FileSystemStorage()
            fs1.save(date1, imageField)
            path1 = fs1.url(date1)
            res=Staff.objects.filter(LOGIN=request.session['lid']).update(name=name,photo=path1,gender=gender,dob=dob,phone=phone,mail=mailid,state=state,housename=house,place=place,district=district,pincode=pincode,qualification=qualification,current_institution=current_inst)
            res1 = Login.objects.filter(id=request.session['lid']).update(username=username)
    else:
        res = Staff.objects.filter(LOGIN=request.session['lid']).update(name=name,gender=gender, dob=dob,
                                                                        phone=phone, mail=mailid,
                                                                        state=state, housename=house, place=place,
                                                                        district=district, pincode=pincode,
                                                                        qualification=qualification,
                                                                        current_institution=current_inst)
        res1 = Login.objects.filter(id=request.session['lid']).update(username=username)
    return HttpResponse( '''<script>alert("Profile Edited!");window.location="/myapp/stf_view_profile/"</script>''')

def stf_edit_questions(request,id):
    res = TestQuestions.objects.get(id=id)
    tname = Test.objects.filter(STAFF__LOGIN_id=request.session['lid'])
    q = Question.objects.filter(LOGIN_id=request.session['lid'])
    return render(request,'staff/edit_questions.html', {'data': res ,'tname': tname ,'q':q})
def stf_edit_questions_post(request):
    id= request.POST['id']
    tq = TestQuestions.objects.get(id=id)
    # tq.TEST_id = request.POST['tname']
    qut = request.POST.get('quest', None)
    qut2 = request.POST.get('question', None)

    if qut:  # If 'quest' is present in the POST data
        tq.question = qut
        tq.type = 'typed'
    elif qut2:  # If 'question' is present in the POST data
        tq.question = Question.objects.get(id=qut2).question
        tq.type = 'list'

    tq.option1 = request.POST['opt1']
    tq.option2 = request.POST['opt2']
    tq.option3 = request.POST['opt3']
    tq.option4 = request.POST['opt4']
    tq.option4 = request.POST['crctans']
    tq.save()
    return HttpResponse('''<script>alert("Test Question Edited.");window.location="/myapp/stf_manage_test_questions/"</script>''')

def stf_edit_test(request,id):
    res=Test.objects.get(id=id)
    return render(request,'staff/edit_test.html',{'data':res})

def stf_edit_test_post(request):
    t_name=request.POST['t_name']
    date=request.POST['date']
    time=request.POST['time']
    id=request.POST['id1']
    obj = Test.objects.get(id=id)
    obj.testname = t_name
    obj.date = date
    obj.time = time
    # obj.STAFF = Staff.objects.get(LOGIN_id=request.session['lid'])
    obj.save()
    return HttpResponse('''<script>alert("Test Edited.");window.location="/myapp/stf_manage_test/"</script>''')

def stf_fileupload(request):
    return render(request,'staff/fileupload.html')

def stf_fileupload_post(request):
    if request.POST["submit"] == "Upload":


        text=request.POST['textarea']

        from myapp.pipelines import pipeline

        qa = pipeline("question-generation")
        s = qa(text)
        print(s)

        from transformers import pipeline
        qa = pipeline("summarization")


        i= text.split('.')
        r=""

        m=""
        for h in i:
            if len(m) <3000:
                m= m+ h
            else:
                r=r+ qa(m)[0]['summary_text']

                m=h


        if len(m) > 0:
            r = r + qa(m)[0]['summary_text']

        print(s)

        m=r
        language = "en"
        max_ngram_size = 1
        deduplication_threshold = 0.9
        numOfKeywords = 5
        custom_kw_extractor = yake.KeywordExtractor(lan=language, n=max_ngram_size,
                                                    dedupLim=deduplication_threshold, top=numOfKeywords,
                                                    features=None)

        keywords = custom_kw_extractor.extract_keywords(text)
        print(keywords)


        qatext=""
        for i in s:
            print(i)
            qatext= qatext+"\n"+i['question']+"\n"+i['answer']+"\n"


        qatext=qatext+"\n SUMMARY \n"+m+"\n"
        request.session["qatext"]=qatext
        print(qatext)

        #save into database
        for j in s:
            q=Question()
            q.question=j['question']
            q.answer=j['answer']
            q.LOGIN_id=request.session['lid']
            q.save()




        return render(request, 'staff/fileupload.html', {'data': s, 'm': m, 'keywords': keywords})

    elif request.POST["submit"] == "Export":

        from datetime import datetime
        date = datetime.now().strftime("%Y%m%d%H%M%S")

        if request.POST["options"]=="voice":
            from gtts import gTTS
            tts = gTTS(request.session["qatext"])
            tts.save("C:\\Users\\ayson\\Downloads\\Telegram Desktop\\EL24-01\\media\\docs\\"+date+".mp3")
            return HttpResponse("<a href='/media/docs/" + date + ".mp3' download>Download</a>")

        else:

            from fpdf import FPDF
            from PyPDF2 import PdfFileMerger

            def create_pdf(text):
                # Create a new FPDF object
                pdf = FPDF()

                # Add a new page to the PDF
                pdf.add_page()

                # Set the font and font size
                pdf.set_font('Arial', size=12)

                # Write the text to the PDF
                pdf.write(5, text)

                # Save the PDF
                pdf.output("C:\\Users\\ayson\\Downloads\\Telegram Desktop\\EL24-01\\media\\docs\\"+date+".pdf")



            print(request.session["qatext"])

            c=request.session["qatext"]

            f=["\n","\t","a","b","c","d","e","f","g","h","i","j","k","l","m","n","o","p","q","r","s","t","u","v","w","x","y","z","A","B","C","D","E","F","G","H","I","J","K","L","M","N","O","P","Q","R","S","T","U","V","W","X","Y","Z",",","."," ","?",":",";","1","2","3","4","5","7","8","9","0","6","(",")"]
            CC=""
            for i in c:
                if i in f:
                    CC= CC+i


            c=CC


            create_pdf(c)
            return HttpResponse("<a href='/media/docs/" + date + ".pdf' download>Download</a>")


    else:
        from datetime import datetime
        file=request.FILES['textfield']
        date=datetime.now().strftime("%Y%m%d%H%M%S")+file.name
        fs = FileSystemStorage()
        fs.save('docs/'+date, file)
        # path = fs.url(date)
        # file.save("C:\\Users\\dell\\PycharmProjects\\easy_learn\\static\\uploaded_file\\"+date)
        path="/media/uploaded_file/"+date

        if '.pdf' in file.name:
            import PyPDF2
            pdfFileObj = open("C:\\Users\\ayson\\Downloads\\Telegram Desktop\\EL24-01\\media\\docs\\"+date, 'rb')
            pdfReader = PyPDF2.PdfReader(pdfFileObj)

            te=""

            for i in pdfReader.pages :
                pageObj = i
                print(pageObj.extract_text())
                text=pageObj.extract_text()
                te= te+" "+ text
            pdfFileObj.close()
            return render(request,'staff/fileupload.html',{'d':te})

        elif '.docx' in file.name:
            import docx
            doc = docx.Document(r"C:\\Users\\ayson\\Downloads\\Telegram Desktop\\EL24-01\\media\\docs\\"+date)
            allText = ""
            for docpara in doc.paragraphs:
                allText= allText+ docpara.text
            print(allText)
            return render(request,'staff/fileupload.html',{'d':allText})


def stf_manage_test(request):
    res=Test.objects.filter(STAFF__LOGIN_id=request.session['lid'])
    return render(request,'staff/manage_test.html',{'data':res})
def stf_manage_test_post(request):
    search=request.POST['search']
    res = Test.objects.filter(STAFF__LOGIN_id=request.session['lid'],testname__icontains=search)
    return render(request, 'staff/manage_test.html', {'data': res})

def stf_manage_tq(request):
    res = Test.objects.filter(STAFF__LOGIN_id=request.session['lid'])
    return render(request,'staff/manage_test_questions.html',{'data':res})
def stf_manage_tq_post(request):
    search = request.POST['tests']
    res = Test.objects.filter(STAFF__LOGIN_id=request.session['lid'])
    quest = TestQuestions.objects.filter(TEST__STAFF__LOGIN_id=request.session['lid'], TEST__testname__icontains=search)
    return render(request, 'staff/manage_test_questions.html', {'quest': quest, 'data' : res})

def stf_send_complaints(request):
    return render(request,'staff/send_complaints.html')
def stf_send_complaints_post(request):
    complaint=request.POST['textarea']
    from datetime import datetime
    date = datetime.now()
    obj=Complaints()
    obj.date=date
    obj.complaint=complaint
    obj.status='pending'
    obj.LOGIN_id=request.session['lid']
    obj.save()
    return HttpResponse('''<script>alert("Complaint Sent Successfully");window.location="/myapp/stf_home/"</script>''')

def stf_signup(request):
    return render(request,'staff/signup.html')
def stf_signup_post(request):
    from datetime import datetime
    date1="stf_photos/"+datetime.now().strftime("%Y%m%d-%H%M%S")+"-1.jpg"
    fs1=FileSystemStorage()
    imageField=request.FILES['imageField']
    fs1.save(date1,imageField)
    path1=fs1.url(date1)
    name=request.POST['textfield']
    gender=request.POST['gender']
    dob=request.POST['textfield1']
    mailid=request.POST['textfield2']
    phone=request.POST['textfield3']
    state=request.POST['textfield4']
    house=request.POST['textfield5']
    place=request.POST['textfield6']
    district=request.POST['textfield7']
    pincode=request.POST['textfield8']
    qualification=request.POST['textfield9']
    current_inst=request.POST['textfield10']
    username =request.POST['textfield11']
    password = request.POST['textfield12']
    confirm_password = request.POST['textfield13']
    # if (type(phone) != int):
    #     return HttpResponse(
    #         '''<script>alert("Phone number should only contain digits.");window.location="/myapp/stf_sign_up/"</script>''')

    date2 = "stf_proof/" + datetime.now().strftime("%Y%m%d-%H%M%S") + "-2.jpg"
    fs2 = FileSystemStorage()
    file = request.FILES['proof']
    fs2.save(date2, file)
    path2 = fs2.url(date2)

    flag = Login.objects.filter(username=username).exists()
    if flag:
        return HttpResponse(
            '''<script>alert("Username unavailable.");window.location="/myapp/stf_sign_up/"</script>''')

    else:
        if (confirm_password == password):
            lobj = Login()
            lobj.type='staff'
            lobj.username=username
            lobj.password=password
            lobj.save()

            sobj=Staff()
            sobj.LOGIN=lobj
            sobj.photo=path1
            sobj.name=name
            sobj.dob=dob
            sobj.gender=gender
            sobj.mail=mailid
            sobj.phone=phone
            sobj.state=state
            sobj.housename=house
            sobj.place=place
            sobj.district=district
            sobj.pincode=pincode
            sobj.qualification=qualification
            sobj.current_institution=current_inst
            sobj.proof=path2
            sobj.status="pending"
            sobj.save()

            return HttpResponse('''<script>alert("Registered successfully. Waiting for confirmation");window.location="/myapp/login/"</script>''')
        else:
            return HttpResponse(
                '''<script>alert("Password and Confirm password must be same.");window.location="/myapp/stf_sign_up/"</script>''')

def stf_vw_profile(request):
    res=Staff.objects.get(LOGIN=request.session['lid'])
    return render(request,'staff/view_profile.html',{'staff':res})

def stf_vw_reply(request):
    res=Complaints.objects.filter(LOGIN_id=request.session['lid'])
    return render(request,'staff/view_reply.html',{'data':res})
def stf_vw_reply_post(request):
    from_date=request.POST['textfield']
    to_date=request.POST['textfield2']
    res = Complaints.objects.filter(date__range=[from_date, to_date],LOGIN_id=request.session['lid'])
    return render(request,'staff/view_reply.html',{'data':res})

def stf_vw_testresult(request,id):
    res=Result.objects.filter(TEST__STAFF__LOGIN_id=request.session['lid'],TEST_id=id)
    return render(request,'staff/view_test_results.html',{'data':res,'id':id})
def stf_vw_testresult_post(request,id):
    search=request.POST['search']
    res = Result.objects.filter(TEST__STAFF__LOGIN_id=request.session['lid'],TEST__testname__icontains=search,TEST_id=id)
    return render(request, 'staff/view_test_results.html', {'data': res,'id':id})

def stf_voice(request):
    return render(request,'staff/voice.html')
def stf_voice_post(request):
    d= request.POST["a"]
    return render(request, 'staff/fileupload.html', {'d': d})



#===================User======================
def usr_home(request):
    return render(request,'user/usr_home.html')

def usr_changepswd(request):
    return render(request,'user/change_passwd.html')
def usr_changepswd_post(request):
    old_password=request.POST['textfield']
    new_password=request.POST['textfield2']
    confirm_password = request.POST['textfield3']
    obj = Login.objects.filter(id=request.session['lid'], password=old_password)
    if obj.exists():
        obj = Login.objects.get(id=request.session['lid'], password=old_password)

        if (new_password == confirm_password):
            obj.password = confirm_password
            obj.save()
            return HttpResponse(
                '''<script>alert("Password Changed Successfully.");window.location="/myapp/login/"</script>''')
        else:
            return HttpResponse(
                '''<script>alert("New and confirm password must be same");window.location="/myapp/usr_change_passwd/"</script>''')
    else:
        return HttpResponse(
            '''<script>alert("Incorrect Password");window.location="/myapp/usr_change_passwd/"</script>''')



def usr_attend_test(request):
    data = Test.objects.all()
    res=[]
    import datetime
    current_time= datetime.datetime.now()
    for i in data:
        flag = not Result.objects.filter(TEST_id=i.id, USER__LOGIN_id=request.session['lid']).exists()
        test_time= datetime.datetime.combine(i.date,i.time)
        time_diff= current_time - test_time
        if time_diff >= datetime.timedelta(minutes=-5):
            if time_diff <= datetime.timedelta(minutes=60):
                if flag:
                    res.append(i)
    return render(request,'user/attend test.html', {'data':res})

def usr_attend_test_post(request):
    name = request.POST['textfield2']
    data = Test.objects.filter(testname__icontains=name)
    res = []
    import datetime
    current_time = datetime.datetime.now()
    for i in data:
        flag = not Result.objects.filter(TEST_id=i.id, USER__LOGIN_id=request.session['lid']).exists()
        test_time = datetime.datetime.combine(i.date, i.time)
        time_diff = current_time - test_time
        if time_diff >= datetime.timedelta(minutes=-5):
            if time_diff <= datetime.timedelta(minutes=60):
                if flag:
                    res.append(i)
    return render(request,'user/attend test.html', {'data':res})

def usr_complaints_agnst_stf(request):
    res = Staff.objects.filter(status='approved')
    return render(request,'user/complaint against staff.html',{'data':res})
def usr_complaints_agnst_stf_post(request):
    complaint = request.POST['textarea']
    from datetime import datetime
    date = datetime.now()
    obj = Staff_Complaints()
    obj.date = date
    obj.complaint = complaint
    obj.status = 'pending'
    obj.USER= User.objects.get(LOGIN_id=request.session['lid'])
    obj.STAFF_id = request.POST['staff']
    obj.save()

    return HttpResponse('''<script>alert("Complaint Sent Successfully");window.location="/myapp/usr_home/"</script>''')

def usr_edit_profile(request):
    res = User.objects.get(LOGIN=request.session['lid'])
    return render(request, 'user/edit profile.html', {'user': res})

def usr_edit_profile_post(request):
    name=request.POST['textfield']
    gender=request.POST['gender']
    dob=request.POST['textfield1']
    username=request.POST['textfield2']
    phone=request.POST['textfield3']
    mailid=request.POST['textfield4']
    place=request.POST['textfield5']
    current_inst=request.POST['textfield6']
    if 'imageField' in request.FILES:
        imageField = request.FILES['imageField']
        if imageField!="":
            from datetime import datetime
            date1 = "usr_photos/" + datetime.now().strftime("%Y%m%d-%H%M%S") + "-1.jpg"
            fs1 = FileSystemStorage()
            fs1.save(date1, imageField)
            path1 = fs1.url(date1)
            res=User.objects.filter(LOGIN=request.session['lid']).update(name=name,photo=path1,gender=gender,dob=dob,phone=phone,mail=mailid,place=place,
                                                                          current_education_status=current_inst)
            res1 = Login.objects.filter(id=request.session['lid']).update(username=username)
    else:
        res = User.objects.filter(LOGIN=request.session['lid']).update(name=name,gender=gender, dob=dob,
                                                                        phone=phone, mail=mailid,
                                                                        place=place,
                                                                        current_education_status=current_inst)
        res1 = Login.objects.filter(id=request.session['lid']).update(username=username)
    return HttpResponse( '''<script>alert("Profile Edited!");window.location="/myapp/usr_view_profile/"</script>''')

def usr_file_upload(request):
    return render(request,'user/fileupload.html')
def usr_file_upload_post(request):
    if  request.POST["submit"]=="Upload":

        text=request.POST['textarea']

        from myapp.pipelines import pipeline

        qa = pipeline("question-generation")
        s = qa(text)
        print(s)

        from transformers import pipeline
        qa = pipeline("summarization")


        i= text.split('.')
        r=""

        m=""
        for h in i:
            if len(m) <3000:
                m= m+ h
            else:
                r=r+ qa(m)[0]['summary_text']

                m=h


        if len(m) >0:
            r = r + qa(m)[0]['summary_text']

        print(s)

        m=r
        language = "en"
        max_ngram_size = 1
        deduplication_threshold = 0.9
        numOfKeywords = 5
        custom_kw_extractor = yake.KeywordExtractor(lan=language, n=max_ngram_size,
                                                    dedupLim=deduplication_threshold, top=numOfKeywords,
                                                    features=None)

        keywords = custom_kw_extractor.extract_keywords(text)
        print(keywords)


        qatext=""
        for i in s:
            print(i)
            qatext= qatext+"\n"+i['question']+"\n"+i['answer']+"\n"


        qatext=qatext+"\n SUMMARY \n"+m+"\n"
        request.session["qatext"]=qatext
        print(qatext)


        return render(request, 'user/fileupload.html', {'data': s, 'm': m, 'keywords': keywords})

    elif request.POST["submit"] == "Export":

        from datetime import datetime
        date = datetime.now().strftime("%Y%m%d%H%M%S")

        if request.POST["options"]=="voice":
            from gtts import gTTS
            tts = gTTS(request.session["qatext"])
            tts.save("C:\\Users\\ayson\\Downloads\\Telegram Desktop\\EL24-01\\media\\docs\\"+date+".mp3")
            return HttpResponse("<a href='/media/docs/" + date + ".mp3' download>Download</a>")

        else:

            from fpdf import FPDF
            from PyPDF2 import PdfFileMerger

            def create_pdf(text):
                # Create a new FPDF object
                pdf = FPDF()

                # Add a new page to the PDF
                pdf.add_page()

                # Set the font and font size
                pdf.set_font('Arial', size=12)

                # Write the text to the PDF
                pdf.write(5, text)

                # Save the PDF
                pdf.output("C:\\Users\\ayson\\Downloads\\Telegram Desktop\\EL24-01\\media\\docs\\"+date+".pdf")



            print(request.session["qatext"])

            c=request.session["qatext"]

            f=["\n","\t","a","b","c","d","e","f","g","h","i","j","k","l","m","n","o","p","q","r","s","t","u","v","w","x","y","z","A","B","C","D","E","F","G","H","I","J","K","L","M","N","O","P","Q","R","S","T","U","V","W","X","Y","Z",",","."," ","?",":",";","1","2","3","4","5","7","8","9","0","6","(",")"]
            CC=""
            for i in c:
                if i in f:
                    CC= CC+i


            c=CC


            create_pdf(c)


            return HttpResponse("<a href='/media/docs/" + date + ".pdf' download>Download</a>")


    else:
        from datetime import datetime
        file=request.FILES['textfield']
        date=datetime.now().strftime("%Y%m%d%H%M%S")+file.name
        fs = FileSystemStorage()
        fs.save('docs/'+date, file)
        # path = fs.url(date)
        # file.save("C:\\Users\\dell\\PycharmProjects\\easy_learn\\static\\uploaded_file\\"+date)
        path="/media/uploaded_file/"+date

        if '.pdf' in file.name:
            import PyPDF2
            pdfFileObj = open("C:\\Users\\ayson\\Downloads\\Telegram Desktop\\EL24-01\\media\\docs\\"+date, 'rb')
            pdfReader = PyPDF2.PdfReader(pdfFileObj)

            te=""

            for i in pdfReader.pages :
                pageObj = i
                print(pageObj.extract_text())
                text=pageObj.extract_text()
                te= te+" "+ text
            pdfFileObj.close()
            return render(request,'user/fileupload.html',{'d':te})

        elif '.docx' in file.name:
            import docx
            doc = docx.Document(r"C:\\Users\\ayson\\Downloads\\Telegram Desktop\\EL24-01\\media\\docs\\"+date)
            allText = ""
            for docpara in doc.paragraphs:
                allText= allText+ docpara.text
            print(allText)
            return render(request,'user/fileupload.html',{'d':allText})


def usr_questions(request):
    obj=Question.objects.filter(LOGIN=request.session['lid'])
    return render(request,'user/questions.html',{'data':obj})
def usr_questions_post(request):
    radio1=request.POST['radio1']
    radio2=request.POST['radio2']
    radio3=request.POST['radio3']
    radio4=request.POST['radio4']
    return HttpResponse("Ok")

def usr_send_complaints(request):
    return render(request,'user/send complaints.html')
def usr_send_complaints_post(request):
    complaint=request.POST['textarea']
    from datetime import datetime
    date = datetime.now()
    obj=Complaints()
    obj.date=date
    obj.complaint=complaint
    obj.status='pending'
    obj.LOGIN_id=request.session['lid']
    obj.save()
    return HttpResponse('''<script>alert("Complaint Sent Successfully");window.location="/myapp/usr_home/"</script>''')


def usr_send_reviews(request):
    return render(request,'user/send reviews.html')
def usr_send_reviews_post(request):
    review=request.POST['review']
    rating = request.POST['rating']
    res = Reviews()
    res.review = review
    res.rating = rating
    from datetime import datetime
    res.date = datetime.now().date()
    res.type = 'user'
    res.LOGIN_id = request.session['lid']
    res.save()
    return HttpResponse('''<script>alert("Review has been sent.");window.location="/myapp/usr_home/"</script>''')



def usr_sign_up(request):
    return render(request,'user/sign up.html')
def usr_sign_up_post(request):
    from datetime import datetime
    date1 = "usr_photos/" + datetime.now().strftime("%Y%m%d-%H%M%S") + "-1.jpg"
    fs1 = FileSystemStorage()
    imageField = request.FILES['imageField']
    fs1.save(date1, imageField)
    name = request.POST['text']
    gender = request.POST['gender']
    dob = request.POST['text1']
    mailid = request.POST['text2']
    phone = request.POST['text3']
    place = request.POST['text4']
    current_inst = request.POST['text5']
    username = request.POST['text6']
    password = request.POST['text7']
    confirm_password = request.POST['text8']

    # if(type(phone)!= int):
    #     return HttpResponse(
    #         '''<script>alert("Phone number should only contain digits.");window.location="/myapp/usr_sign_up/"</script>''')

    path2 = fs1.url(date1)
    flag = Login.objects.filter(username=username).exists()
    if flag:
        return HttpResponse(
            '''<script>alert("Username unavailable.");window.location="/myapp/usr_sign_up/"</script>''')

    else:
        if (confirm_password == password):
            obj2 = Login()
            obj2.username = username
            obj2.password = password
            obj2.type = "user"
            obj2.save()

            obj = User()
            obj.LOGIN_id=obj2.id
            obj.name=name
            obj.gender=gender
            obj.dob=dob
            obj.photo=path2
            obj.mail=mailid
            obj.phone= phone
            obj.place=place
            obj.current_education_status = current_inst
            obj.save()

            return HttpResponse('''<script>alert("Registered successfully.");window.location="/myapp/login/"</script>''')

        else:
            return HttpResponse('''<script>alert("Password and Confirm password must be same.");window.location="/myapp/usr_sign_up/"</script>''')

def usr_vw_profile(request):
    res = User.objects.get(LOGIN_id=request.session['lid'])
    return render(request, 'user/view profile.html', {'user': res})

def usr_vw_reply(request):
    res = Complaints.objects.filter(LOGIN_id=request.session['lid'])
    return render(request, 'user/view reply.html', {'data': res})
def usr_vw_reply_post(request):
    from_date = request.POST['textfield']
    to_date = request.POST['textfield2']
    res = Complaints.objects.filter(date__range=[from_date, to_date], LOGIN_id=request.session['lid'])
    return render(request, 'user/view reply.html', {'data': res})


def usr_vw_result(request):
    res = Result.objects.filter(USER__LOGIN_id=request.session['lid'])
    return render(request,'user/view result.html', {'data': res})
def usr_vw_result_post(request):
    search=request.POST['textfield']
    res = Result.objects.filter(TEST__testname__icontains=search,USER__LOGIN_id=request.session['lid'])
    return render(request,'user/view result.html', {'data': res})

def usr_exam_center(request,id):
    obj = TestQuestions.objects.filter(TEST_id=id)
    return render(request, 'user/usr_exam_center.html', {'data':obj,'id':id})
def usr_exam_center_post(request):
    id=request.POST['id']
    obj= TestQuestions.objects.filter(TEST_id=id)
    mrk=0
    j=0
    for i in obj:
        j=j+1
        ans= "ans" + str(j)
        if request.POST[ans] == i.correct_answer:
            mrk=mrk+1
    mobj=Result()
    print(id)
    mobj.TEST_id=id
    mobj.USER=User.objects.get(LOGIN__id=request.session['lid'])
    mobj.mark=mrk
    mobj.save()
    return HttpResponse('''<script>alert("Test has been submitted.");window.location="/myapp/usr_attend_test/"</script>''')

def usr_voice(request):
    return render(request,'user/voice.html')
def usr_voice_post(request):
    d= request.POST["a"]
    return render(request, 'user/fileupload.html', {'d': d})
